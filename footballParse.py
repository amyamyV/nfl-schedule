import requests
import os
import re
import calendar
from docxtpl import DocxTemplate
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.dml.color import ColorFormat
from docx.shared import RGBColor
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Length
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from docx.enum.text import WD_LINE_SPACING

#using Python3. To run: python3 _fileName_
#to install stuff: pip3 install --user _packageName_
#to open: open _fileName_

#method to insert horizontal line

#convert month names to number

def getFootballFormatted(string_of_week_url_espn):
    abbr_to_num = {name: num for num, name in enumerate(calendar.month_abbr) if num}

    url = string_of_week_url_espn
    response = requests.get(url)

    soup = BeautifulSoup(response.content, 'lxml')
    #print(soup.prettify())

    games=[]  # a list to store games

    table = soup.find('div', attrs = {'id':'sched-container'})

    dates = table.findAll('h2', attrs = {'class':'table-caption'})
    daygamesdiv = table.findAll('div', attrs = {'class':'responsive-table-wrap'})

    #if new: document = Document()  instead

    #with lines (may need adjusting)
    document = DocxTemplate("line_template.docx")
    #for new doc
    #document = Document()
    #font = document.styles['Normal'].font
    #10 font, Verdana, BOLD, dark blue
    #document.font.name = 'Verdana'
    #document.font.bold = True
    #document.add_heading('Weekly Football Schedule', 0)
    text = ''
    daycount = 0;
    gamecount = 1;

    for daygames in daygamesdiv:
        formatted_date_text = ''
        m = re.match("(\w+),\s(\w+)\s(\d+)", dates[daycount].text)
        m.group(1)
        m.group(2)
        m.group(3)
        monthNumber = str(abbr_to_num[m.group(2)[:3]])
        formatted_date_text += monthNumber + '-' + m.group(3) + ' ' + m.group(1)[:3]
        daycount += 1
        datetext= document.add_paragraph()
        datetext.paragraph_format.space_after = Pt(0)
        datetext_run = datetext.add_run(formatted_date_text.upper())
        datetext_run.font.color.rgb = RGBColor(0x0, 0x0, 0x0)
        datetext_run.font.size = Pt(8)
        datetext_run.left_margin = Inches(.2)
        datetext_run.font.name = 'Verdana'
        datetext_run.bold = True
        text = ''
        for game in daygames.findAll('a', attrs = {'class':'team-name'}):
            text += game.span.text + '\n'
            if (gamecount % 2) == 0:
                     text += '\n'
                     #paragraph_format.space_after = Pt(0)
            gamecount += 1
        text = text[:-2]
        paragraph = document.add_paragraph()
        paragraph_run = paragraph.add_run(text.upper())
        paragraph_run.font.color.rgb = RGBColor(0x17, 0x24, 0x92)
        paragraph_run.font.size = Pt(10)
        paragraph_run.left_margin = Inches(.2)
        paragraph_run.font.name = 'Verdana'
        paragraph_run.bold = True
        paragraph.paragraph_format.space_before = Pt(0)

    document.save('weeklyFootball.docx')
    os.system('open weeklyFootball.docx')
